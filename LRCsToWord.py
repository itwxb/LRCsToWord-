import os
from docx import Document

# 获取当前脚本所在的目录
current_directory = os.path.dirname(os.path.abspath(__file__))

# 构建“歌词”文件夹的路径
lyrics_folder_path = os.path.join(current_directory, "歌词")

# 创建一个新的Word文档
doc = Document()

# 遍历“歌词”文件夹中的所有文件
for filename in os.listdir(lyrics_folder_path):
    if filename.endswith(".lrc"):
        # 构建LRC文件的完整路径
        file_path = os.path.join(lyrics_folder_path, filename)

        # 读取LRC文件的内容
        with open(file_path, "r", encoding="utf-8") as file:
            lrc_content = file.read()

        # 将LRC内容添加到Word文档中
        # 可以根据需要添加标题或分隔线
        doc.add_heading(f"Lyrics for {filename}", level=2)
        doc.add_paragraph(lrc_content)
        # 如果不需要在每个LRC文件内容后添加额外的分隔，可以注释掉或删除下一行
        # doc.add_paragraph('')  # 添加一个空段落作为分隔

# 保存Word文档到当前目录下（或您指定的其他路径）
output_path = os.path.join(current_directory, "compiled_lyrics.docx")
doc.save(output_path)

print(f"Word document saved at {output_path}")
